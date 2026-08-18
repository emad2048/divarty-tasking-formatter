"""Assemble a full Weekly Tasking Order from ontology index + LLM bodies."""

from __future__ import annotations

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH

from divarty_tasking_formatter.classification import add_classification_banner
from divarty_tasking_formatter.models import WeeklyOrderData
from divarty_tasking_formatter.render import (
    create_document,
    render_summary_block,
    render_tasking_detail,
    _add_paragraph,
    _add_run,
    _apply_indent,
)


def format_weekly_order(
    order: WeeklyOrderData,
    template_path: str | None = None,
    *,
    nlt_dates_value_bold: bool = False,
) -> tuple[DocumentType, list[str]]:
    """Letterhead, references, task index, each tasking body, one signature."""
    warnings: list[str] = []
    doc, used_template = create_document(template_path)
    add_classification_banner(doc, level=order.classification)
    _render_letterhead(doc, order, used_template)
    _render_references(doc, order, used_template)

    taskings = sorted(
        order.taskings,
        key=lambda t: (t.task_number is None, t.task_number or 0),
    )
    for tasking in taskings:
        warnings.extend(tasking.warnings)
        render_summary_block(doc, tasking, nlt_dates_value_bold, used_template)
    for tasking in taskings:
        render_tasking_detail(doc, tasking, used_template, warnings)

    _render_signature(doc, order, used_template)
    return doc, warnings


def _p(doc: DocumentType, used_template: bool, style_key: str = "section"):
    paragraph = _add_paragraph(doc, style_key, used_template)
    _apply_indent(paragraph, "section")
    return paragraph


def _render_letterhead(doc: DocumentType, order: WeeklyOrderData, used_template: bool) -> None:
    hq = _p(doc, used_template)
    hq.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(hq, order.headquarters, bold=True)

    loc = _p(doc, used_template)
    _add_run(loc, order.location, bold=True)

    date_p = _p(doc, used_template)
    _add_run(date_p, order.date_line, bold=True)

    title = _p(doc, used_template)
    _add_run(
        title,
        f"1CD DIVARTY WEEKLY TASKING ORDER {order.order_number}",
        bold=True,
    )


def _render_references(doc: DocumentType, order: WeeklyOrderData, used_template: bool) -> None:
    heading = _p(doc, used_template)
    _add_run(heading, "REFERENCES:", bold=True)
    if order.references:
        for ref in order.references:
            line = _p(doc, used_template)
            _add_run(line, ref, bold=False)
    else:
        omitted = _p(doc, used_template)
        _add_run(omitted, "Omitted", bold=False)

    tz = _p(doc, used_template)
    _add_run(tz, "Time Zone Used Throughout Order: ", bold=True)
    _add_run(tz, order.time_zone, bold=False)

    org = _p(doc, used_template)
    _add_run(org, "Task Organization: ", bold=True)
    _add_run(org, order.task_organization, bold=False)


def _render_signature(doc: DocumentType, order: WeeklyOrderData, used_template: bool) -> None:
    ack = _p(doc, used_template)
    _add_run(ack, "ACKNOWLEDGE:", bold=True)
    if order.acknowledge_name:
        name = _p(doc, used_template)
        _add_run(name, order.acknowledge_name.upper(), bold=False)
    rank = _p(doc, used_template)
    _add_run(rank, order.acknowledge_rank, bold=False)

    official = _p(doc, used_template)
    _add_run(official, "OFFICIAL:", bold=True)
    if order.official_name:
        s3 = _p(doc, used_template)
        _add_run(s3, order.official_name.upper(), bold=False)
    title = _p(doc, used_template)
    _add_run(title, order.official_title, bold=False)

    if order.attachments:
        att_h = _p(doc, used_template)
        _add_run(att_h, "ATTACHMENTS:", bold=True)
        for index, name in enumerate(order.attachments, start=1):
            line = _p(doc, used_template)
            _add_run(line, f"Enclosure {index}: {name}", bold=False)
