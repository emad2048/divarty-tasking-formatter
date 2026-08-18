"""Render TaskingData to a python-docx document chunk.

Formatting follows 1CD DIVARTY WTO 26-42-001 unit convention:
``1. / a. / (1) / i.`` — not ATP 5-0.1 ``1. / a. / (1) / (a)``.

TODO(verify): a unit .docx template with letterhead/named styles has not been
provided. Named-style lookup is attempted when a template is loaded; otherwise
this module applies the fallback font/indent spec via inline runs.
"""

from __future__ import annotations

import logging
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from divarty_tasking_formatter.classification import add_classification_banner
from divarty_tasking_formatter.models import OutlineNode, TaskingData, TextSpan

logger = logging.getLogger(__name__)

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
MARGIN_INCHES = 1.0
SPACE_AFTER_PT = 6
WARNING_COLOR = RGBColor(0xFF, 0x00, 0x00)

# Fallback indents (inches). Hanging indent is first_line_indent = -hang.
# Level 2 hang is sized for markers like "(10) ".
LEVEL_LEFT_IN = {
    "section": 0.0,
    "subparagraph": 0.25,
    "unit_header": 0.25,
    "item0": 0.5,   # (1)
    "item1": 0.75,  # i.
    "item2": 1.0,   # (a) fallback
    "item3": 1.25,  # (i) fallback
}
LEVEL_HANG_IN = {
    "section": 0.30,
    "subparagraph": 0.25,
    "unit_header": 0.0,
    "item0": 0.40,
    "item1": 0.30,
    "item2": 0.40,
    "item3": 0.40,
}

# TODO(verify): replace these placeholder names with the template's actual
# paragraph style IDs once the .docx template is provided. When a name exists
# on doc.styles it is applied; mixed bold/underline still uses runs.
TEMPLATE_STYLE_NAMES = {
    "section": "TaskingLevel0",
    "subparagraph": "TaskingLevel1",
    "unit_header": "TaskingUnitHeader",
    "item0": "TaskingLevel2",
    "item1": "TaskingLevel3",
    "item2": "TaskingLevel4",
    "item3": "TaskingLevel5",
    "summary": "TaskingSummary",
    "body_header": "TaskingBodyHeader",
}

_ROMAN_MAP = (
    (1000, "m"),
    (900, "cm"),
    (500, "d"),
    (400, "cd"),
    (100, "c"),
    (90, "xc"),
    (50, "l"),
    (40, "xl"),
    (10, "x"),
    (9, "ix"),
    (5, "v"),
    (4, "iv"),
    (1, "i"),
)


def format_tasking(
    data: TaskingData,
    template_path: str | None = None,
    *,
    nlt_dates_value_bold: bool = False,
    include_summary: bool = True,
) -> tuple[DocumentType, list[str]]:
    """Build a docx chunk for one tasking.

    ``nlt_dates_value_bold`` defaults to False (plain NLT value, matching
    WTO 26-42-001). This is a minor style choice, not a hard requirement.

    Returns the document plus warnings from parse and render.
    """
    warnings = list(data.warnings)
    doc, used_template = create_document(template_path)
    add_classification_banner(doc, level="UNCLASSIFIED")
    if include_summary:
        render_summary_block(doc, data, nlt_dates_value_bold, used_template)
    render_tasking_detail(doc, data, used_template, warnings)
    return doc, warnings


def create_document(template_path: str | None = None) -> tuple[DocumentType, bool]:
    """Load the unit template if given; otherwise a blank TNR 12pt document.

    A 26-42-001 .docx (letterhead / header-footer / named styles) should be
    passed as template_path when available. Body outline is still drawn here.
    """
    used_template = bool(template_path)
    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()
        _apply_fallback_defaults(doc)
        _remove_placeholder_paragraph(doc)
    return doc, used_template


def item_marker(level: int, number: int) -> str:
    """Army outline marker for a numbered item. Level 0 = (1), 1 = i., then fallbacks."""
    if level <= 0:
        return f"({number})"
    if level == 1:
        return f"{_to_roman(number)}."
    if level == 2:
        return f"({_to_alpha(number)})"
    if level == 3:
        return f"({_to_roman(number)})"
    return f"{number}."


def _to_roman(n: int) -> str:
    n = max(int(n), 1)
    parts: list[str] = []
    remaining = n
    for value, symbol in _ROMAN_MAP:
        while remaining >= value:
            parts.append(symbol)
            remaining -= value
    return "".join(parts)


def _to_alpha(n: int) -> str:
    """1 → a, 26 → z, 27 → aa."""
    n = max(int(n), 1)
    chars: list[str] = []
    while n > 0:
        n, rem = divmod(n - 1, 26)
        chars.append(chr(ord("a") + rem))
    return "".join(reversed(chars))


def _apply_fallback_defaults(doc: DocumentType) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_PT)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(SPACE_AFTER_PT)
    pf.space_before = Pt(0)
    pf.first_line_indent = Inches(0)

    for section in doc.sections:
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)


def _remove_placeholder_paragraph(doc: DocumentType) -> None:
    if len(doc.paragraphs) == 1 and doc.paragraphs[0].text == "":
        element = doc.paragraphs[0]._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _style_exists(doc: DocumentType, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _add_paragraph(doc: DocumentType, style_key: str, used_template: bool) -> Paragraph:
    paragraph = doc.add_paragraph()
    # Prefer template named styles when present; otherwise (and as a supplement)
    # apply fallback left/hanging indent below.
    if used_template:
        style_name = TEMPLATE_STYLE_NAMES.get(style_key)
        if style_name and _style_exists(doc, style_name):
            paragraph.style = doc.styles[style_name]
            return paragraph
    _apply_indent(paragraph, style_key)
    return paragraph


def _apply_indent(paragraph: Paragraph, style_key: str) -> None:
    left = LEVEL_LEFT_IN.get(style_key, 0.0)
    hang = LEVEL_HANG_IN.get(style_key, 0.25)
    pf = paragraph.paragraph_format
    pf.left_indent = Inches(left)
    pf.first_line_indent = Inches(-hang) if hang else Inches(0)
    pf.space_after = Pt(SPACE_AFTER_PT)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _style_run(run: Run, *, bold: bool | None = False, underline: bool | None = False) -> None:
    run.bold = bool(bold)
    run.underline = bool(underline)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_run(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    underline: bool = False,
) -> Run:
    run = paragraph.add_run(text)
    _style_run(run, bold=bold, underline=underline)
    return run


def _add_spans(
    paragraph: Paragraph,
    spans: list[TextSpan],
    *,
    force_bold: bool = False,
    underline: bool = False,
) -> None:
    for span in spans:
        if not span.text:
            continue
        _add_run(
            paragraph,
            span.text,
            bold=force_bold or span.bold,
            underline=underline,
        )


def _body_text_preview(node: OutlineNode, limit: int = 40) -> str:
    text = "".join(span.text for span in node.body).strip()
    if len(text) > limit:
        return text[:limit] + "..."
    return text or (node.title or "")


def _note_warning(warnings: list[str], message: str) -> None:
    warnings.append(message)
    logger.warning(message)


def _add_warning_run(paragraph: Paragraph, message: str) -> None:
    run = paragraph.add_run(" " + message)
    _style_run(run, bold=False, underline=False)
    run.font.color.rgb = WARNING_COLOR


def render_summary_block(
    doc: DocumentType,
    data: TaskingData,
    nlt_dates_value_bold: bool,
    used_template: bool,
) -> None:
    number = data.task_number if data.task_number is not None else "?"
    if data.designation:
        label = f"Task #{number} ({data.designation}): "
    else:
        label = f"Task #{number}: "
    title = (data.title or "").upper()

    line = _add_paragraph(doc, "summary", used_template)
    _apply_indent(line, "section")
    _add_run(line, label, bold=True)
    _add_run(line, title, bold=False)

    affected = _add_paragraph(doc, "summary", used_template)
    _apply_indent(affected, "section")
    _add_run(affected, "Affected Unit/Staff: ", bold=True)
    _add_run(affected, data.affected_unit_staff, bold=False)

    nlt = _add_paragraph(doc, "summary", used_template)
    _apply_indent(nlt, "section")
    _add_run(nlt, "NLT Dates: ", bold=True)
    _add_run(nlt, data.nlt_dates, bold=nlt_dates_value_bold)

    bluf = _add_paragraph(doc, "summary", used_template)
    _apply_indent(bluf, "section")
    _add_run(bluf, "BLUF: ", bold=True)
    _add_run(bluf, data.bluf, bold=False)


def render_tasking_detail(
    doc: DocumentType,
    data: TaskingData,
    used_template: bool,
    warnings: list[str],
) -> None:
    """TASK # header plus 1/2/3 body (no index/summary block)."""
    _render_body_header(doc, data, used_template)
    for index, section in enumerate(data.sections):
        _render_node(
            doc,
            section,
            section_number=index + 1,
            parent_kind=None,
            item_level=0,
            item_number=1,
            used_template=used_template,
            warnings=warnings,
        )


def _render_body_header(doc: DocumentType, data: TaskingData, used_template: bool) -> None:
    number = data.task_number if data.task_number is not None else "?"
    if data.designation:
        text = f"TASK #{number} ({data.designation}): {(data.title or '').upper()}"
    else:
        text = f"TASK #{number}: {(data.title or '').upper()}"
    paragraph = _add_paragraph(doc, "body_header", used_template)
    _apply_indent(paragraph, "section")
    _add_run(paragraph, text, bold=True, underline=True)


def _title_with_period(title: str) -> str:
    title = title.strip()
    if not title.endswith("."):
        title += "."
    return title


def _render_node(
    doc: DocumentType,
    node: OutlineNode,
    *,
    section_number: int,
    parent_kind: str | None,
    item_level: int,
    item_number: int,
    used_template: bool,
    warnings: list[str],
    sub_letter_index: int = 0,
) -> None:
    if node.kind == "section":
        _render_section(doc, node, section_number, used_template, warnings)
        return
    if node.kind == "subparagraph":
        _render_subparagraph(
            doc, node, sub_letter_index, used_template, warnings
        )
        return
    if node.kind == "unit_header":
        _render_unit_header(doc, node, used_template, warnings)
        return
    if node.kind == "item":
        _render_item(
            doc,
            node,
            parent_kind=parent_kind,
            item_level=item_level,
            item_number=item_number,
            used_template=used_template,
            warnings=warnings,
        )


def _render_children(
    doc: DocumentType,
    parent: OutlineNode,
    *,
    used_template: bool,
    warnings: list[str],
    child_item_level: int,
) -> None:
    letter_index = 0
    item_number = 0
    for child in parent.children:
        if child.kind == "subparagraph":
            _render_node(
                doc,
                child,
                section_number=0,
                parent_kind=parent.kind,
                item_level=0,
                item_number=1,
                used_template=used_template,
                warnings=warnings,
                sub_letter_index=letter_index,
            )
            letter_index += 1
        elif child.kind == "item":
            item_number += 1
            _render_node(
                doc,
                child,
                section_number=0,
                parent_kind=parent.kind,
                item_level=child_item_level,
                item_number=item_number,
                used_template=used_template,
                warnings=warnings,
            )
        else:
            _render_node(
                doc,
                child,
                section_number=0,
                parent_kind=parent.kind,
                item_level=child_item_level,
                item_number=1,
                used_template=used_template,
                warnings=warnings,
            )


def _render_section(
    doc: DocumentType,
    node: OutlineNode,
    number: int,
    used_template: bool,
    warnings: list[str],
) -> None:
    paragraph = _add_paragraph(doc, "section", used_template)
    _apply_indent(paragraph, "section")
    _add_run(paragraph, f"{number}. ", bold=False, underline=False)
    title = _title_with_period(node.title or "SECTION")
    _add_run(paragraph, title, bold=True, underline=True)
    if node.omitted:
        _add_run(paragraph, " Omitted", bold=False, underline=False)
    elif node.body:
        _add_run(paragraph, " ", bold=False, underline=False)
        _add_spans(paragraph, node.body)

    _render_children(
        doc,
        node,
        used_template=used_template,
        warnings=warnings,
        child_item_level=0,
    )


def _render_subparagraph(
    doc: DocumentType,
    node: OutlineNode,
    letter_index: int,
    used_template: bool,
    warnings: list[str],
) -> None:
    paragraph = _add_paragraph(doc, "subparagraph", used_template)
    _apply_indent(paragraph, "subparagraph")
    letter = _to_alpha(letter_index + 1)
    label = f"{letter}. {_title_with_period(node.title or '')}"
    _add_run(paragraph, label, bold=False, underline=True)
    if node.omitted:
        _add_run(paragraph, " Omitted", bold=False, underline=False)
    elif node.body:
        _add_run(paragraph, " ", bold=False, underline=False)
        _add_spans(paragraph, node.body)

    _render_children(
        doc,
        node,
        used_template=used_template,
        warnings=warnings,
        child_item_level=0,
    )


def _render_unit_header(
    doc: DocumentType,
    node: OutlineNode,
    used_template: bool,
    warnings: list[str],
) -> None:
    paragraph = _add_paragraph(doc, "unit_header", used_template)
    _apply_indent(paragraph, "unit_header")
    name = (node.title or "").rstrip(":")
    _add_run(paragraph, f"{name}:", bold=True, underline=False)
    if node.body:
        _add_run(paragraph, " ", bold=False)
        _add_spans(paragraph, node.body)
    _render_children(
        doc,
        node,
        used_template=used_template,
        warnings=warnings,
        child_item_level=0,
    )


def _item_style_key(level: int) -> str:
    if level <= 0:
        return "item0"
    if level == 1:
        return "item1"
    if level == 2:
        return "item2"
    return "item3"


def _render_item(
    doc: DocumentType,
    node: OutlineNode,
    *,
    parent_kind: str | None,
    item_level: int,
    item_number: int,
    used_template: bool,
    warnings: list[str],
) -> None:
    flagged = False
    if parent_kind == "section":
        flagged = True
    if node.list_depth is not None and node.list_depth != item_level:
        flagged = True
    if item_level >= 2:
        flagged = True

    style_key = _item_style_key(item_level)
    paragraph = _add_paragraph(doc, style_key, used_template)
    _apply_indent(paragraph, style_key)
    marker = item_marker(item_level, item_number)
    force_bold = item_level == 1  # roman (level 3 in the outline table) is bold
    _add_run(paragraph, f"{marker} ", bold=force_bold, underline=False)
    _add_spans(paragraph, node.body, force_bold=force_bold)

    if item_level >= 2:
        _note_warning(
            warnings,
            "Numbered list deeper than verified unit levels 0-3; "
            f"using fallback marker near {_body_text_preview(node)!r}.",
        )
    if flagged:
        preview = _body_text_preview(node)
        message = f'[FORMATTING WARNING: unexpected nesting near "{preview}"]'
        _note_warning(warnings, message)
        _add_warning_run(paragraph, message)

    nested_number = 0
    for child in node.children:
        if child.kind == "item":
            nested_number += 1
            _render_item(
                doc,
                child,
                parent_kind="item",
                item_level=item_level + 1,
                item_number=nested_number,
                used_template=used_template,
                warnings=warnings,
            )
        else:
            _render_node(
                doc,
                child,
                section_number=0,
                parent_kind="item",
                item_level=item_level + 1,
                item_number=1,
                used_template=used_template,
                warnings=warnings,
            )
