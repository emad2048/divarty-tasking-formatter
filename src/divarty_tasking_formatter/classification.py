"""Classification markings for each page (WTO 26-42-001: UNCLASSIFIED header/footer).

TODO(verify): banner color by classification level, portion markings, and
distribution statements are not specified. This implementation is black,
centered, bold Times New Roman — matching the example PDF's UNCLASSIFIED
banners, not a color-coded DoD template.
"""

from __future__ import annotations

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def add_classification_banner(doc: DocumentType, level: str = "UNCLASSIFIED") -> None:
    """Centered bold classification line in header and footer of every section."""
    label = (level or "UNCLASSIFIED").strip().upper() or "UNCLASSIFIED"
    for section in doc.sections:
        _write_banner(section.header, label)
        _write_banner(section.footer, label)


def _write_banner(container, label: str) -> None:
    paragraph = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Drop leftover placeholder runs from a blank document.
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(label)
    run.bold = True
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
