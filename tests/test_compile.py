from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx.text.paragraph import Paragraph

from divarty_tasking_formatter.compile import format_weekly_order
from divarty_tasking_formatter.foundry import generate_weekly_order_bytes
from divarty_tasking_formatter.models import WeeklyOrderData
from divarty_tasking_formatter.ontology import merge_orders_and_bodies
from divarty_tasking_formatter.template import write_letterhead_template

FIXTURES = Path(__file__).parent / "fixtures"


def _left_inches(paragraph: Paragraph) -> float:
    value = paragraph.paragraph_format.left_indent
    return 0.0 if value is None else round(float(value.inches), 4)


def _find(paragraphs: list[Paragraph], contains: str) -> Paragraph:
    for paragraph in paragraphs:
        if contains in paragraph.text:
            return paragraph
    raise AssertionError(f"No paragraph contains {contains!r}")


def _merged_order() -> WeeklyOrderData:
    orders = json.loads((FIXTURES / "hhq_orders.json").read_text(encoding="utf-8"))
    estimates = json.loads((FIXTURES / "running_estimates.json").read_text(encoding="utf-8"))
    bodies = json.loads((FIXTURES / "llm_bodies.json").read_text(encoding="utf-8"))
    taskings = merge_orders_and_bodies(orders, estimates, bodies)
    return WeeklyOrderData(
        order_number="26-42-001",
        date_line="16JUL26",
        references=["1CD ORDER 26L4-110 (IIIAC CHANGE OF RESPONSIBILITY)."],
        taskings=taskings,
        acknowledge_name="DVONCH",
        official_name="EANS",
        attachments=["APPOINTMENT ORDERS"],
    )


def test_weekly_order_matches_pdf_labels_and_indents() -> None:
    document, warnings = format_weekly_order(_merged_order())
    texts = [p.text for p in document.paragraphs]
    joined = "\n".join(texts)

    assert "HQ, 1CD DIVARTY" in joined
    assert "FORT HOOD, TX" in joined
    assert "1CD DIVARTY WEEKLY TASKING ORDER 26-42-001" in joined
    assert "REFERENCES:" in joined
    assert "Time Zone Used Throughout Order:" in joined
    assert "Task Organization:" in joined
    assert "(U)" not in joined
    assert "None." not in joined
    assert "1-82 FA:" not in joined
    assert "PART 1" not in joined

    summary = _find(document.paragraphs, "Task #1 (FORAC):")
    assert summary.runs[0].bold
    assert "FRAGORD 1" in summary.text.upper()

    unit = _find(document.paragraphs, "3-16 FAR:")
    assert unit.text.strip() == "3-16 FAR:"
    assert unit.runs[0].bold
    assert _left_inches(unit) == pytest.approx(0.25, abs=0.01)

    item = _find(document.paragraphs, "(1) Provide salute battery")
    assert _left_inches(item) == pytest.approx(0.5, abs=0.01)

    honor = _find(document.paragraphs, "Honor Guard:")
    assert any(run.bold and "Honor Guard:" in run.text for run in honor.runs)

    roman = _find(document.paragraphs, "i. Sub-item requiring roman")
    assert _left_inches(roman) == pytest.approx(0.75, abs=0.01)
    assert all(run.bold for run in roman.runs if run.text.strip())

    command = _find(document.paragraphs, "a. Command.")
    assert command.text.endswith("Omitted")

    inform = _find(document.paragraphs, "Task #3 (INFORM):")
    assert "SEPARATION" in inform.text.upper()
    all_header = _find(document.paragraphs, "All:")
    assert all_header.runs[0].bold

    assert any("ACKNOWLEDGE:" in t for t in texts)
    assert any(t.strip() == "DVONCH" for t in texts)
    assert any("OFFICIAL:" in t for t in texts)
    assert any("Enclosure 1: APPOINTMENT ORDERS" in t for t in texts)

    header_text = document.sections[0].header.paragraphs[0].text
    footer_text = document.sections[0].footer.paragraphs[0].text
    assert "UNCLASSIFIED" in header_text
    assert "UNCLASSIFIED" in footer_text
    assert not any("FORMATTING WARNING" in w for w in warnings)


def test_generate_weekly_order_bytes_and_template(tmp_path: Path) -> None:
    template = write_letterhead_template(tmp_path / "letterhead.docx")
    orders = json.loads((FIXTURES / "hhq_orders.json").read_text(encoding="utf-8"))
    estimates = json.loads((FIXTURES / "running_estimates.json").read_text(encoding="utf-8"))
    bodies = (FIXTURES / "llm_bodies.json").read_text(encoding="utf-8")
    payload, warnings = generate_weekly_order_bytes(
        orders,
        estimates,
        bodies,
        order_number="26-42-001",
        date_line="16JUL26",
        template_path=str(template),
        acknowledge_name="DVONCH",
        official_name="EANS",
    )
    assert payload[:2] == b"PK"
    assert isinstance(warnings, list)
