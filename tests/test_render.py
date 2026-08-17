from __future__ import annotations

import pytest
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from divarty_tasking_formatter.parse import parse_input
from divarty_tasking_formatter.render import format_tasking


def _left_inches(paragraph: Paragraph) -> float:
    value = paragraph.paragraph_format.left_indent
    return 0.0 if value is None else round(float(value.inches), 4)


def _is_bold(run) -> bool:
    return bool(run.bold)


def _is_underline(run) -> bool:
    return bool(run.underline)


def _find(paragraphs: list[Paragraph], contains: str) -> Paragraph:
    for paragraph in paragraphs:
        if contains in paragraph.text:
            return paragraph
    raise AssertionError(f"No paragraph contains {contains!r}")


def _format_fixture(fixture_text, name: str):
    data = parse_input(fixture_text(name))
    document, warnings = format_tasking(data)
    return document, warnings, data


def test_ceremony_summary_and_body_header_runs(fixture_text) -> None:
    document, warnings, _ = _format_fixture(fixture_text, "ceremony_units.md")
    paragraphs = document.paragraphs

    summary = paragraphs[0]
    assert summary.text.startswith("Task #1 (FORAC):")
    assert "FRAGORD 1 TO IIIAC CHANGE OF RESPONSIBILITY" in summary.text
    assert _is_bold(summary.runs[0])
    assert "Task #1 (FORAC):" in summary.runs[0].text
    assert not _is_bold(summary.runs[1])

    affected = _find(paragraphs, "Affected Unit/Staff:")
    assert _is_bold(affected.runs[0])
    assert affected.runs[0].text.startswith("Affected Unit/Staff:")
    assert not _is_bold(affected.runs[1])

    nlt = _find(paragraphs, "NLT Dates:")
    assert _is_bold(nlt.runs[0])
    assert not _is_bold(nlt.runs[1])

    header = _find(paragraphs, "TASK #1 (FORAC):")
    assert header.text.startswith("TASK #1 (FORAC):")
    assert _is_bold(header.runs[0])
    assert _is_underline(header.runs[0])
    assert _left_inches(header) == pytest.approx(0.0, abs=0.01)


def test_ceremony_outline_numbering_indents_and_unit_headers(fixture_text) -> None:
    document, _, _ = _format_fixture(fixture_text, "ceremony_units.md")
    paragraphs = document.paragraphs

    situation = _find(paragraphs, "1. SITUATION.")
    assert situation.runs[0].text == "1. "
    assert not _is_bold(situation.runs[0])
    assert not _is_underline(situation.runs[0])
    assert situation.runs[1].text.startswith("SITUATION.")
    assert _is_bold(situation.runs[1])
    assert _is_underline(situation.runs[1])
    assert _left_inches(situation) == pytest.approx(0.0, abs=0.01)

    concept = _find(paragraphs, "a. Concept of Operations.")
    assert _is_underline(concept.runs[0])
    assert not _is_bold(concept.runs[0])
    assert _left_inches(concept) == pytest.approx(0.25, abs=0.01)
    following = "".join(run.text for run in concept.runs[1:])
    assert "1st Cavalry Division" in following
    assert not any(_is_underline(run) for run in concept.runs[1:] if run.text.strip())

    unit = _find(paragraphs, "3-16 FAR:")
    assert unit.text.strip() == "3-16 FAR:"
    assert _is_bold(unit.runs[0])
    assert _left_inches(unit) == pytest.approx(0.25, abs=0.01)

    hhbn = _find(paragraphs, "HHBN:")
    assert _is_bold(hhbn.runs[0])
    assert _left_inches(hhbn) == pytest.approx(0.25, abs=0.01)

    item = _find(paragraphs, "(1) Provide salute battery")
    assert item.text.startswith("(1) ")
    assert _left_inches(item) == pytest.approx(0.5, abs=0.01)
    assert not _is_bold(item.runs[0])

    honor = _find(paragraphs, "Honor Guard:")
    bold_runs = [run.text for run in honor.runs if _is_bold(run)]
    assert any("Honor Guard:" in text for text in bold_runs)

    roman = _find(paragraphs, "i. Sub-item requiring roman")
    assert roman.text.startswith("i. ")
    assert _left_inches(roman) == pytest.approx(0.75, abs=0.01)
    assert all(_is_bold(run) for run in roman.runs if run.text.strip())


def test_omitted_labels_are_kept(fixture_text) -> None:
    document, _, _ = _format_fixture(fixture_text, "omitted_sections.md")
    texts = [p.text for p in document.paragraphs]
    assert any("c. Coordinating Instructions." in t and t.endswith("Omitted") for t in texts)
    assert any("a. Command." in t and t.endswith("Omitted") for t in texts)
    assert any("a. Concept of Operations." in t for t in texts)
    assert any("b. Tasks to Subordinate Units." in t for t in texts)
    omitted_coord = _find(document.paragraphs, "c. Coordinating Instructions.")
    assert omitted_coord.text == "c. Coordinating Instructions. Omitted"
    omitted_command = _find(document.paragraphs, "a. Command.")
    assert omitted_command.text == "a. Command. Omitted"


def test_inform_all_unit_and_concept_title(fixture_text) -> None:
    document, _, _ = _format_fixture(fixture_text, "inform_all.md")
    summary = document.paragraphs[0]
    assert summary.text.startswith("Task #3 (INFORM):")
    assert any("a. Concept of the Operation." in p.text for p in document.paragraphs)
    all_header = _find(document.paragraphs, "All:")
    assert all_header.text.strip() == "All:"
    assert _is_bold(all_header.runs[0])
    assert _left_inches(all_header) == pytest.approx(0.25, abs=0.01)
    command = _find(document.paragraphs, "a. Command.")
    assert command.text.endswith("Omitted")


def test_malformed_nesting_warns_and_still_renders(fixture_text) -> None:
    document, warnings, _ = _format_fixture(fixture_text, "malformed_nesting.md")
    assert any("FORMATTING WARNING" in w for w in warnings)
    flagged = [p for p in document.paragraphs if "FORMATTING WARNING" in p.text]
    assert flagged
    assert "no parent subparagraph" in flagged[0].text
    assert any(
        run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
        for run in flagged[0].runs
        if run.font.color and run.font.color.rgb
    )
    assert any("1. SITUATION." in p.text for p in document.paragraphs)
